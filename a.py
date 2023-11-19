import urllib
import warnings
from pathlib import Path as p
import re
import pandas as pd
from langchain import PromptTemplate
from langchain.chains.summarize import load_summarize_chain
from langchain.document_loaders import PyPDFLoader
from langchain.text_splitter import CharacterTextSplitter
from langchain.document_loaders import PyPDFLoader
# from .parser import Retrive
import os
import streamlit as st
import io
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import pandas as pd
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
import os
import openai
from openai import OpenAI
from langchain.embeddings.openai import OpenAIEmbeddings
from langchain.text_splitter import CharacterTextSplitter
from langchain.vectorstores import ElasticVectorSearch, Pinecone, Weaviate, FAISS

from langchain.chains.question_answering import load_qa_chain
from langchain.llms import OpenAI
from langchain.prompts import PromptTemplate


#os.environ["OPENAI_API_KEY"] = "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"
os.environ["OPENAI_API_KEY"] = "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"
api_key = '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'



embeddings = OpenAIEmbeddings()
text_splitter = CharacterTextSplitter(
  separator="\n",
  chunk_size=1000,
  chunk_overlap=200,
  length_function=len,
)





def ask_gpt(text, question):
    """
    Asks a question from the provided text using GPT-3.5.

    Args:
    text (str): The extracted text from the PDF.
    question (str): The question to ask.

    Returns:
    str: The answer from GPT-3.5.
    """
    #openai.api_key = api_key
    client = OpenAI()
    texts = text_splitter.split_text(text)
    docsearch = FAISS.from_texts(texts, embeddings)
    docs = docsearch.similarity_search(question)
    print(len(docs))
    response = openai.chat.completions.create(
        model="gpt-3.5-turbo",
        messages=[
            {
                "role": "system",
                "content": "You are a helpful assistant. From the provided docs answer the queries in newly seperated lines."
            },
            {
                "role": "user",
                "content": f"Text: {docs}"
            },
            {
                "role": "user",
                "content": f"Question: {question}"
            }
        ]
    )
    print(response)

    return response


class Retrive:
  def __init__(self, filepath):
    loader = PyPDFLoader(filepath)
    pages = loader.load_and_split()
    text_splitter = CharacterTextSplitter(chunk_size=1000, chunk_overlap=0)
    context = "\n\n".join(str(p.page_content) for p in pages)
    texts = text_splitter.split_text(context)

    st = ""
    for i in texts:
      st += i
    self.st=st

    question = "From the given text find out land address, land tax, settlement date, land status and plan detials."
    answer = ask_gpt(st, question)
    answer=answer.choices[0].message.content.split('\n')
    #dict_keys(['Land address', 'Land tax', 'Settlement date', 'Land status', 'Plan details'])
    property_info = {}
    for item in k:
        key, value = map(str.strip, item.split(':', 1))
        property_info[key] = value
    sents = st.split("\n")
    keys = ['vendor ',"purchaser",'improvements','inclusion','exclusions']

    dc = {}
    for item in keys:
        for x in sents:
            if item in x:
                print(item, x)
                dc[item] = [sents[sents.index(x)], sents[sents.index(x) + 1], sents[sents.index(x) + 2], sents[sents.index(x) + 3], sents[sents.index(x) + 4]]
                break  # Stop iterating after finding the first occurrence

# Now, dc will only contain values for keys based on the first occurrence in sentences.


    dc['vendor '] = dc['vendor '][0] #vendor's name
    dc['Settlement date'] = property_info["Settlement date"] #date of completion
    dc['Land status'] = property_info['Land status']
    dc['improvements'] = dc['improvements'][0:2]
    #dc['land (address']=dc['land (address'][2:4]
    dc['exclusions']=dc['exclusions'][0]
    dc['Land tax']=dc['Land tax'][0]

    self.dc = dc
    self.property_info=property_info


  def getName(self):
    return self.dc['vendor '].split(' ',1)[1]

  def getLandaddress(self):
    return self.property_info['Land address']

  def getPlandetails(self):
    return self.property_info['Plan details']

  def getSettlementdate(self):
    doc = self.property_info['Settlement date']
    if len(doc) == 0:
      return "Need to be confirmed"
    return doc[0].split('date for completion')[1]

  def getLandstatus(self):
    sts = re.findall(r'[\uf0fe☒]\s*([^☐\uf0fe]+)', self.dc['VACANT'][0])
    if len(sts) == 0:
      return "Further clarification is also required of whether the Vendor will be able to provide vacant possession on settlement."
    return sts

  def getPrice(self):
    pattern = re.compile(r'\bbalance\b\s*(\$[\d,.]+)')
    matches = pattern.findall(self.st)
    print(matches)
    #price = dc['purchaser’s solicitor    '].split('balance')[1].strip()
    if len(matches) == 0:
      return "TBA"
    return matches[0]

  def getImprovments(self):
    selected_options = re.findall(r'(?:X|\uf0fe)\s*([^X\s]+)', self.dc['improvements'][0])

    if len(selected_options) == 0:
      return "Need to be confirmed"
    elif len(selected_options) == 1:
      return selected_options[0]
    else:
      impvs = ""
      for i in range(0, len(selected_options)-1):
        impvs += selected_options[i] + ','
      impvs += 'and' + selected_options[-1]
      return impvs

  def getInclusions(self):
    inc = re.findall(r'\uf0fe\s*([^☐\uf0fe]+)', self.dc['inclusion'][0])

    if len(inc) != 0:
      return "Inclusions are marked under the inclusion tab of the contract"
    return "Inclusions are not marked under the inclusion tab of the contract"

  def getExclusions(self):
    exc = self.dc['exclusions'].split('exclusions')[1]
    return exc

  def getLandtax(self):
    ltx = self.property_info[ 'Land tax']
    if len(ltx) == 0:
      return "Land tax is not marked as adjustable or not adjustable"
    elif ltx[0].strip().lower() == 'no':
      return "Land tax is marked as not adjustable"
    else:
      return "Land tax is marked as adjustable"


def save_uploaded_file(uploaded_file, folder_path):
    with open(os.path.join(folder_path, uploaded_file.name), "wb") as f:
        f.write(uploaded_file.getbuffer())



# ... (previous code)

def main():
    st.title("Information retriever")
    uploaded_file = st.file_uploader("Choose a PDF file", type=["pdf"])

    if uploaded_file is not None:
        save_uploaded_file(uploaded_file, "Data/")
        pdf_file = os.path.join("Data/", uploaded_file.name)

    
    if st.button("Retrieve"):
      try:
        retriever = Retrive(pdf_file)
      except:
        st.write("Error: Cannot load the file")
      # Create a DataFrame to store the data
      data = {
            "Field": [
              "Field",
                "Settlement Date",
                "Land Status",
                "Price",
                "Improvements",
                "Inclusions",
                "Exclusions",
                "Land Tax"
            ],
            "Value": [
              "Value",
                retriever.getSettlementdate(),
                retriever.getLandstatus(),
                retriever.getPrice(),
                retriever.getImprovments(),
                retriever.getInclusions(),
                retriever.getExclusions(),
                retriever.getLandtax()
            ]
      }

      df = pd.DataFrame(data)

      # Save data to Excel file
      excel_file = "output.xlsx"
      df.to_excel(excel_file, index=False, header=False)

      # Create a Word document
      doc = Document()
      doc.add_heading("Information Retrieved from PDF", 0)
      doc.add_paragraph("Vendor's Name: " + retriever.getName())
      doc.add_paragraph("Land Address: " + retriever.getLandaddress())
      doc.add_paragraph("Plan Details: " + retriever.getPlandetails())

      # Read the Excel file and add data to the Word document
      df = pd.read_excel(excel_file)
      table = doc.add_table(rows=0, cols=2)
      table.style = 'Table Grid'
      table.alignment = WD_ALIGN_VERTICAL.CENTER

      for index, row in df.iterrows():
          row_cells = table.add_row().cells
          row_cells[0].text = row["Field"]
          row_cells[1].text = str(row["Value"])

      # Save the Word document
      output_docx_file = str(uploaded_file.name).replace(".pdf", ".docx")  
      doc.save(output_docx_file)

      bio = io.BytesIO()
      doc.save(bio)
      # Provide a link to download the generated Word document
      # st.markdown(f"**[Download Word Document](/{output_docx_file})**")
      st.download_button(
          label="Download Word Document",
          data=bio.getvalue(),
          file_name=output_docx_file,
          mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

if __name__ == "__main__":
    main()
